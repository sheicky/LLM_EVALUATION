from openai import OpenAI
from sentence_transformers import SentenceTransformer 
from sklearn.metrics.pairwise import cosine_similarity
from pinecone import Pinecone
from langchain_pinecone import PineconeVectorStore
from langchain.docstore.document import Document
from langchain.embeddings import OpenAIEmbeddings
from langchain_community.embeddings import HuggingFaceEmbeddings
import tempfile
import os 
from dotenv import load_dotenv
import google.generativeai as genai
import streamlit as st
import io
from pypdf import PdfReader
import docx
from datetime import datetime


load_dotenv()



class PerformRag : 
    """
    This class is used to perform RAG on the user input.
    """


    def connect_to_pinecone(self) : 
        pinecone_api_key = os.getenv("PINECONE_API_KEY")
        pc = Pinecone(api_key=pinecone_api_key) 
        pinecone_index = pc.Index("llm-eval")
        return pinecone_index



    def embedding(self,text,model_name="sentence-transformers/all-mpnet-base-v2") : 
        model = SentenceTransformer(model_name) 
        return model.encode(text)
    
    def get_file_content(self, file_content):
        try:
            # Pour les fichiers déjà en bytes (uploadés via Streamlit)
            if isinstance(file_content, bytes):
                return file_content.decode('utf-8', errors='ignore')
            
            # Pour les fichiers sur le disque
            elif isinstance(file_content, str):
                with open(file_content, "r", encoding='utf-8', errors='ignore') as file:
                    return file.read()
            
            return str(file_content)
        except Exception as e:
            st.error(f"Error reading file: {str(e)}")
            return None
    
    def parse_document(self, file_content, file_name):
        """Parse different document types"""
        try:
            file_extension = file_name.lower().split('.')[-1]
            file_stream = io.BytesIO(file_content)
            
            if file_extension == 'pdf':
                pdf_reader = PdfReader(file_stream)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text.strip()
                
            elif file_extension in ['docx', 'doc']:
                doc = docx.Document(file_stream)
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                return text.strip()
                
            elif file_extension == 'txt':
                return file_content.decode('utf-8', errors='ignore').strip()
                
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
                
        except Exception as e:
            st.error(f"Error parsing document: {str(e)}")
            return None

    def get_documents(self, file_content, file_name=None):
        if not file_content:
            return []
        
        try:
            # Parser le document selon son type
            if file_name:
                parsed_content = self.parse_document(file_content, file_name)
                if not parsed_content:
                    st.error("Failed to parse document")
                    return []
    
                
                # Créer le document avec métadonnées détaillées
                doc = Document(
                    page_content=parsed_content,
                    metadata={
                        "source": file_name,
                        "type": file_name.split('.')[-1],
                        "length": len(parsed_content),
                        "timestamp": str(datetime.now())
                    }
                )
            else:
                # Pour le texte brut
                doc = Document(
                    page_content=str(file_content),
                    metadata={"source": "direct_input"}
                )
            
            return [doc]
            
        except Exception as e:
            st.error(f"Error processing document: {str(e)}")
            return []


    def get_vector_store(self, documents, embedding_model=None):
        # Créer un objet embedding
        if embedding_model is None:
            embedding_model = HuggingFaceEmbeddings(
                model_name="sentence-transformers/all-mpnet-base-v2"
            )
        
        vector_store = PineconeVectorStore.from_documents(
            documents=documents,
            embedding=embedding_model,
            index_name="llm-eval",
            namespace="document_attached"
        )
        return vector_store

    #def perform_rag(self,query) : 
    #    pass 


    #def query_pinecone(self,database) : 
    #    pass 




class RagModel(PerformRag) : 

    def __init__(self,top_matches,query) : 
        self.groq_api_key = os.getenv("GROQ_API_KEY") 
        self.gemini_api_key = os.getenv("GEMINI_API_KEY")
        self.system_prompt = """You are an AI assistant. Analyze the provided document and answer questions about its content.
                            Be specific and use information from the document to support your answers."""
        
    
        
        # Modification de la récupération du contexte
        self.contexts = []
        for match in top_matches['matches']:
            # Vérifier toutes les clés disponibles dans metadata
    
            
            # Essayer différentes clés possibles pour le contenu
            content = None
            if 'metadata' in match:
                metadata = match['metadata']
                if 'page_content' in metadata:
                    content = metadata['page_content']
                elif 'text' in metadata:
                    content = metadata['text']
                elif 'content' in metadata:
                    content = metadata['content']
            
            if content:
                self.contexts.append(content)
        
        # Vérifier si nous avons récupéré du contexte
        if not self.contexts:
            st.error("No context was extracted from the matches")
            self.contexts = ["No context available"]
        
        self.augmented_query = (
            "Document Content:\n" + 
            "\n---\n".join(self.contexts) + 
            "\n\nQuestion: " + query
        )
        
        # Debug - Afficher la requête augmentée
        
        
        self.generation_config = {
            "temperature": 0.7,  # Réduit pour plus de précision
            "top_p": 0.95,
            "top_k": 5,
            "max_output_tokens": 8192,
        }


    def gemini_model_response(self, model_name, api_key): 
        if not api_key:
            raise ValueError("GEMINI_API_KEY not found in environment variables")
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(model_name=model_name, generation_config=self.generation_config)
        
        # Formater le prompt pour inclure le contexte système et la requête
        prompt = f"{self.system_prompt}\n\n{self.augmented_query}"
        
        response = model.generate_content(prompt)
        return response.text
    



    def gpt_model_response(self,model_name, api_key) : 
        if not api_key:
            raise ValueError("GROQ_API_KEY not found in environment variables")
        client = OpenAI(
            api_key=api_key,
            base_url="https://api.groq.com/openai/v1"
        )
        llm_response = client.chat.completions.create(
                model=model_name,
                messages=[
                    {"role": "system", "content": self.system_prompt},
                    {"role": "user", "content": self.augmented_query}
                ]
            )

        return llm_response.choices[0].message.content
    


